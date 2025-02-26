# document_processing/parsers/docx_parser.py
import docx
from pathlib import Path
import logging

logger = logging.getLogger(__name__)

class DocxParser:
    """Parser for DOCX documents"""
    
    def __init__(self):
        self.supported_extensions = ['.docx', '.doc']
        
    def can_parse(self, file_path):
        """Check if the file can be parsed with this parser"""
        return Path(file_path).suffix.lower() in self.supported_extensions
    
    def extract_text(self, file_path):
        """Extract text from a DOCX file"""
        try:
            doc = docx.Document(file_path)
            
            # Extract document properties
            metadata = {
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
                "subject": doc.core_properties.subject or "",
                "page_count": len(doc.sections),
            }
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            # Combine all text
            text = "\n".join(paragraphs)
            
            return {
                "text": text,
                "metadata": metadata,
                "tables": tables,
                "file_path": file_path,
                "file_name": Path(file_path).name
            }
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            raise